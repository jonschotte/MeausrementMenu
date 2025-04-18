import streamlit as st
import pandas as pd
import openai
import plotly.express as px
from io import StringIO
import docx
import PyPDF2
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from dotenv import load_dotenv
import os
import time

# Load environment variables
load_dotenv()

# Load your final measurement dataset
df = pd.read_csv("Final_Measurement_Framework.csv")

# Streamlit UI setup
st.set_page_config(page_title="Measurement Finder", layout="wide")
st.title("Measurement Menu")

# OpenAI API Key setup
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    st.warning("⚠️ No OpenAI API key found in environment variables.")
    api_key = st.text_input("Please enter your OpenAI API Key:", type="password", help="Enter your OpenAI API key to enable the measurement analysis features.")
    if not api_key:
        st.error("❌ OpenAI API key is required to use this application. Please enter your API key to continue.")
        st.stop()
    else:
        st.success("✅ API key received successfully!")
        openai.api_key = api_key
else:
    openai.api_key = api_key

# Input method selection
input_method = st.radio(
    "Choose input method:",
    ["Manual Input", "RFP Document Upload"]
)

# Initialize session state for selected metrics if it doesn't exist
if 'selected_metrics' not in st.session_state:
    st.session_state.selected_metrics = []
if 'show_document' not in st.session_state:
    st.session_state.show_document = False
if 'matched_biz_obj' not in st.session_state:
    st.session_state.matched_biz_obj = None
if 'matched_camp_obj' not in st.session_state:
    st.session_state.matched_camp_obj = None
if 'document_path' not in st.session_state:
    st.session_state.document_path = None
if 'relevant_methods' not in st.session_state:
    st.session_state.relevant_methods = None
if 'objectives_matched' not in st.session_state:
    st.session_state.objectives_matched = False
if 'business_context' not in st.session_state:
    st.session_state.business_context = None

def match_objectives(biz_obj, camp_obj):
    # First, get unique business objectives
    unique_biz_objectives = df['Business Objective'].unique()
    
    # Step 1: Match business objective
    biz_prompt = f"""
    Given this business objective:
    {biz_obj}

    Please match it to the most similar business objective from this list:
    {unique_biz_objectives.tolist()}

    Return only the exact match from the list.
    """
    
    try:
        biz_response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an objective matching expert. Match the given business objective to the most similar one from the provided list. Return only the exact match from the list."},
                {"role": "user", "content": biz_prompt}
            ]
        )
        
        matched_biz_obj = biz_response.choices[0].message.content.strip()
        
        # Step 2: Get campaign objectives for the matched business objective
        available_camp_objectives = df[df['Business Objective'] == matched_biz_obj]['Campaign Objective'].unique()
        
        # Match campaign objective
        camp_prompt = f"""
    Given this campaign objective:
    {camp_obj}

    Please match it to the most similar campaign objective from this list (which are all related to the business objective '{matched_biz_obj}'):
    {available_camp_objectives.tolist()}

    Return only the exact match from the list.
    """
        
        camp_response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an objective matching expert. Match the given campaign objective to the most similar one from the provided list. Return only the exact match from the list."},
                {"role": "user", "content": camp_prompt}
            ]
        )
        
        matched_camp_obj = camp_response.choices[0].message.content.strip()
        
        return matched_biz_obj, matched_camp_obj
    except Exception as e:
        st.error(f"Error matching objectives: {str(e)}")
        return None, None

def generate_document():
    try:
        # Create a new Word document
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading('Measurement Plan', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add business context section
        doc.add_heading('Business Context', level=1)
        doc.add_paragraph(st.session_state.business_context)
        
        # Add objectives section
        doc.add_heading('Business and Campaign Objectives', level=1)
        doc.add_paragraph(f"**Business Objective:** {st.session_state.matched_biz_obj}")
        doc.add_paragraph(f"**Campaign Objective:** {st.session_state.matched_camp_obj}")
        
        # Add selected metrics section
        doc.add_heading('Selected Measurement Methods', level=1)
        
        # Get detailed information for each selected metric
        for metric in st.session_state.selected_metrics:
            metric_row = df[df["Measurement Method"] == metric].iloc[0]
            
            # Add metric heading
            doc.add_heading(metric, level=2)
            
            # Add metric details
            doc.add_paragraph(f"**Description:** {metric_row['Description']}")
            doc.add_paragraph(f"**Implementation Cost:** {metric_row['Implementation Cost (1=Low, 5=High)']}")
            doc.add_paragraph(f"**Impact Duration:** {metric_row['Impact Duration (1=Short, 5=Long)']}")
            
            # Generate detailed analysis using OpenAI
            try:
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are a measurement strategy expert. Provide detailed analysis of how this measurement method applies to the business and campaign objectives."},
                        {"role": "user", "content": f"""
                        Business Objective: {st.session_state.matched_biz_obj}
                        Campaign Objective: {st.session_state.matched_camp_obj}
                        Measurement Method: {metric}
                        Description: {metric_row['Description']}
                        
                        Please provide:
                        1. How this measurement directly supports the objectives
                        2. Key considerations for implementation
                        3. Expected impact and value
                        4. Potential challenges and mitigation strategies
                        """}
                    ]
                )
                
                analysis = response.choices[0].message.content
                doc.add_paragraph(analysis)
                
            except Exception as e:
                doc.add_paragraph(f"Error generating detailed analysis: {str(e)}")
            
            doc.add_paragraph("---")
        
        # Save the document with a unique filename
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        filename = f"Measurement_Plan_{timestamp}.docx"
        doc.save(filename)
        st.session_state.document_path = filename
        return True
    except Exception as e:
        st.error(f"Error generating document: {str(e)}")
        return False

def extract_text_from_document(uploaded_file):
    """Extract text from uploaded document based on its type."""
    if uploaded_file.type == "text/plain":
        return uploaded_file.getvalue().decode("utf-8")
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif uploaded_file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        return "\n".join([page.extract_text() for page in pdf_reader.pages])
    return ""

def chunk_text(text, max_chunk_size=4000):
    """Split text into chunks of approximately max_chunk_size characters."""
    words = text.split()
    chunks = []
    current_chunk = []
    current_size = 0
    
    for word in words:
        if current_size + len(word) + 1 > max_chunk_size:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_size = len(word)
        else:
            current_chunk.append(word)
            current_size += len(word) + 1
    
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    
    return chunks

def extract_objectives_from_chunk(chunk):
    """Extract objectives from a text chunk using OpenAI."""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Extract the business objective and campaign objective from the following text. If you find objectives, format the response as 'Business Objective: [objective]' and 'Campaign Objective: [objective]'. If no objectives are found, return 'No objectives found.'"},
                {"role": "user", "content": chunk}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error processing chunk: {str(e)}")
        return "No objectives found."

def process_document(uploaded_file):
    """Process uploaded document and extract objectives."""
    try:
        # Extract text from document
        content = extract_text_from_document(uploaded_file)
        
        # Split content into chunks
        chunks = chunk_text(content)
        
        # Process each chunk
        all_objectives = []
        for i, chunk in enumerate(chunks):
            with st.spinner(f"Processing document part {i+1}/{len(chunks)}..."):
                objectives = extract_objectives_from_chunk(chunk)
                if "No objectives found" not in objectives:
                    all_objectives.append(objectives)
        
        # Combine and deduplicate objectives
        if all_objectives:
            # Use OpenAI to combine and deduplicate objectives
            combined_prompt = "Combine and deduplicate these extracted objectives, keeping only the most relevant ones:\n\n" + "\n\n".join(all_objectives)
            final_response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Combine and deduplicate the objectives, keeping only the most relevant ones. Return exactly one business objective and one campaign objective."},
                    {"role": "user", "content": combined_prompt}
                ]
            )
            return final_response.choices[0].message.content
        else:
            return "No objectives found in the document."
            
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        return None

if input_method == "Manual Input":
    # User input
    business_context = st.text_area("Describe your **business context** (e.g., industry, size, target audience, key challenges)", 
                                  help="This information will help provide better context for your measurement plan")
    biz_obj = st.text_input("Describe your **business objective**")
    camp_obj = st.text_input("Describe your **campaign objective**")

    # Run matching if inputs are provided and not already matched
    if biz_obj and camp_obj and not st.session_state.objectives_matched:
        with st.spinner("Matching objectives..."):
            matched_biz_obj, matched_camp_obj = match_objectives(biz_obj, camp_obj)
            
            if matched_biz_obj and matched_camp_obj:
                # Store matched objectives and relevant methods in session state
                st.session_state.matched_biz_obj = matched_biz_obj
                st.session_state.matched_camp_obj = matched_camp_obj
                st.session_state.business_context = business_context  # Store business context
                st.session_state.relevant_methods = df[
                    (df['Business Objective'] == matched_biz_obj) & 
                    (df['Campaign Objective'] == matched_camp_obj)
                ]['Measurement Method'].unique().tolist()
                st.session_state.objectives_matched = True

else:
    # RFP Document Upload
    uploaded_file = st.file_uploader("Upload RFP Document", type=['txt', 'docx', 'pdf'])
    
    if uploaded_file and not st.session_state.objectives_matched:
        objectives_text = process_document(uploaded_file)
        
        if objectives_text and "No objectives found" not in objectives_text:
            try:
                # Parse the objectives
                biz_obj = objectives_text.split("Business Objective: ")[1].split("\n")[0]
                camp_obj = objectives_text.split("Campaign Objective: ")[1].strip()
                
                st.write("Extracted Objectives:")
                st.write(f"**Business Objective:** {biz_obj}")
                st.write(f"**Campaign Objective:** {camp_obj}")
                
                # Match the extracted objectives
                with st.spinner("Matching objectives..."):
                    matched_biz_obj, matched_camp_obj = match_objectives(biz_obj, camp_obj)
                    
                    if matched_biz_obj and matched_camp_obj:
                        # Store matched objectives and relevant methods in session state
                        st.session_state.matched_biz_obj = matched_biz_obj
                        st.session_state.matched_camp_obj = matched_camp_obj
                        st.session_state.relevant_methods = df[
                            (df['Business Objective'] == matched_biz_obj) & 
                            (df['Campaign Objective'] == matched_camp_obj)
                        ]['Measurement Method'].unique().tolist()
                        st.session_state.objectives_matched = True
            except Exception as e:
                st.error(f"Error parsing objectives: {str(e)}")
        else:
            st.warning("No objectives found in the document. Please try uploading a different document or use manual input.")

# Display results if we have matched objectives
if st.session_state.objectives_matched:
    # Display explanation of matches
    st.markdown("---")
    st.subheader("🔍 Why These Measurements Were Recommended")
    st.markdown(f"These measurements were selected because they are specifically designed for the matched objectives:")
    st.markdown(f"- **Business Objective:** {st.session_state.matched_biz_obj}")
    st.markdown(f"- **Campaign Objective:** {st.session_state.matched_camp_obj}")

    if not st.session_state.relevant_methods:
        st.warning("⚠️ No recommended measurements found for these objectives. Please check the other available metrics below.")
    else:
        st.markdown("---")
        st.subheader("🧪 Recommended Measurements")

        # Create tile layout (4 columns per row)
        cols = st.columns(4)
        for i, method in enumerate(st.session_state.relevant_methods):
            col = cols[i % 4]
            
            with col:
                # Add checkbox for metric selection
                is_selected = st.checkbox(method, key=f"checkbox_{method}", value=method in st.session_state.selected_metrics)
                if is_selected and method not in st.session_state.selected_metrics:
                    st.session_state.selected_metrics.append(method)
                elif not is_selected and method in st.session_state.selected_metrics:
                    st.session_state.selected_metrics.remove(method)
                
                if st.button(method, key=method, help="Click for details"):
                    selected_row = df[df["Measurement Method"] == method].iloc[0]
                    st.markdown("---")
                    st.markdown(f"### {method}")
                    st.markdown(f"**Description:** {selected_row['Description']}")
                    st.markdown(f"**Implementation Cost:** {selected_row['Implementation Cost (1=Low, 5=High)']}")
                    st.markdown(f"**Impact Duration:** {selected_row['Impact Duration (1=Short, 5=Long)']}")

        # Show matrix visualization for recommended metrics
        st.markdown("---")
        st.subheader("📊 Implementation Cost vs Impact Duration Matrix")
        matrix_df = df[df["Measurement Method"].isin(st.session_state.relevant_methods)].copy()
        if not matrix_df.empty:
            fig = px.scatter(matrix_df, 
                            x="Implementation Cost (1=Low, 5=High)", 
                            y="Impact Duration (1=Short, 5=Long)",
                            text="Measurement Method",
                            title="Recommended Measurements Matrix")
            fig.update_traces(textposition='top center', 
                             marker=dict(size=15, line=dict(width=3, color='black')),
                             textfont=dict(size=12, color='black'))
            fig.update_layout(
                plot_bgcolor='white',
                paper_bgcolor='white',
                xaxis=dict(gridcolor='lightgray', linewidth=2),
                yaxis=dict(gridcolor='lightgray', linewidth=2),
                showlegend=False
            )
            st.plotly_chart(fig, use_container_width=True)

    # Add section for other metrics
    st.markdown("---")
    st.subheader("📋 Other Available Metrics")
    other_methods = df[~df["Measurement Method"].isin(st.session_state.relevant_methods)]["Measurement Method"].unique()
    
    # Create tile layout for other metrics (4 columns per row)
    cols = st.columns(4)
    for i, method in enumerate(other_methods):
        col = cols[i % 4]
        
        with col:
            # Add checkbox for metric selection
            is_selected = st.checkbox(method, key=f"other_checkbox_{method}", value=method in st.session_state.selected_metrics)
            if is_selected and method not in st.session_state.selected_metrics:
                st.session_state.selected_metrics.append(method)
            elif not is_selected and method in st.session_state.selected_metrics:
                st.session_state.selected_metrics.remove(method)
            
            if st.button(method, key=f"other_button_{method}", help="Click for details"):
                selected_row = df[df["Measurement Method"] == method].iloc[0]
                st.markdown("---")
                st.markdown(f"### {method}")
                st.markdown(f"**Description:** {selected_row['Description']}")
                st.markdown(f"**Implementation Cost:** {selected_row['Implementation Cost (1=Low, 5=High)']}")
                st.markdown(f"**Impact Duration:** {selected_row['Impact Duration (1=Short, 5=Long)']}")

    # Show selected metrics and generate document button
    if st.session_state.selected_metrics:
        st.markdown("---")
        st.subheader("📝 Selected Metrics for Measurement Plan")
        st.markdown("The following metrics have been selected for your measurement plan:")
        for metric in st.session_state.selected_metrics:
            st.markdown(f"- {metric}")
        
        col1, col2 = st.columns([1, 4])
        with col1:
            if st.button("Generate Measurement Plan Document"):
                with st.spinner("Generating your measurement plan document..."):
                    if generate_document():
                        st.success("✅ Measurement plan document generated successfully!")
                        st.session_state.show_document = True
        
        if st.session_state.show_document and st.session_state.document_path:
            try:
                with open(st.session_state.document_path, "rb") as file:
                    st.download_button(
                        label="📄 Download Measurement Plan",
                        data=file,
                        file_name="Measurement_Plan.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"Error preparing download: {str(e)}")
                st.session_state.show_document = False
                st.session_state.document_path = None
